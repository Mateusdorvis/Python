from docx import Document
from docx.shared import Inches #importa ele , para remendisinar a imagem na página

doc = Document()
doc.add_heading('Título com imagem', level=2)

doc.add_picture(image_path_or_stream=r'C:\xampp\htdocs\progr\python\ARTE\skylite\oi.jpg',width=Inches(3))#add a imagem na página

doc.add_paragraph('Este paragrafo é do python')
#salva o documento
doc.save('documento.docx')
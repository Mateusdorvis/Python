from docx import Document

document = Document()#chama a classe que cria um documento

#ad o titulo na página com tamanho 1
document.add_heading('Este arquivo foi feito pelo python !', level=1)
#add um paragrafo
document.add_paragraph('Este é o novo paragrafo')
#salva o documento
document.save('meu_documento.docx')
#confirma o salvamento
print('Arquivo docx criado')